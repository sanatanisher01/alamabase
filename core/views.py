from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from .models import Questionnaire, Question, Answer, ReferenceDocument
from .ai_service import AIService
from docx import Document
from docx.shared import Pt, RGBColor

def landing_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    return render(request, 'landing.html')

def signup_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('dashboard')
    else:
        form = UserCreationForm()
    return render(request, 'signup.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(data=request.POST)
        if form.is_valid():
            login(request, form.get_user())
            return redirect('dashboard')
    else:
        form = AuthenticationForm()
    return render(request, 'login.html', {'form': form})

def logout_view(request):
    logout(request)
    return redirect('landing')

@login_required
def dashboard_view(request):
    try:
        questionnaires = Questionnaire.objects.filter(user=request.user).order_by('-uploaded_at')
        documents = ReferenceDocument.objects.filter(user=request.user).order_by('-uploaded_at')
        return render(request, 'dashboard.html', {'questionnaires': questionnaires, 'documents': documents})
    except Exception as e:
        print(f"Dashboard error: {e}")
        return render(request, 'dashboard.html', {'questionnaires': [], 'documents': []})

@login_required
def upload_document_view(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        file = request.FILES.get('file')
        doc = ReferenceDocument.objects.create(user=request.user, title=title, file=file)
        ai_service = AIService()
        doc.content = ai_service.extract_text_from_file(doc.file.path)
        doc.save()
        messages.success(request, 'Reference document uploaded!')
        return redirect('dashboard')
    return render(request, 'upload_document.html')

@login_required
def upload_questionnaire_view(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        file = request.FILES.get('file')
        questionnaire = Questionnaire.objects.create(user=request.user, title=title, file=file)
        messages.success(request, 'Questionnaire uploaded!')
        return redirect('process_questionnaire', pk=questionnaire.pk)
    return render(request, 'upload_questionnaire.html')

@login_required
def process_questionnaire_view(request, pk):
    questionnaire = get_object_or_404(Questionnaire, pk=pk, user=request.user)
    if not questionnaire.processed:
        ai_service = AIService()
        text = ai_service.extract_text_from_file(questionnaire.file.path)
        questions = ai_service.parse_questions(text)
        for idx, q_text in enumerate(questions, 1):
            Question.objects.create(questionnaire=questionnaire, text=q_text, order=idx)
        questionnaire.processed = True
        questionnaire.save()
        messages.success(request, f'Extracted {len(questions)} questions!')
    return redirect('review_answers', pk=questionnaire.pk)

@login_required
def generate_answers_view(request, pk):
    questionnaire = get_object_or_404(Questionnaire, pk=pk, user=request.user)
    documents = ReferenceDocument.objects.filter(user=request.user)
    doc_list = [{'title': doc.title, 'content': doc.content} for doc in documents]
    ai_service = AIService()
    
    Answer.objects.filter(question__questionnaire=questionnaire).delete()
    
    for question in questionnaire.questions.all():
        chunks = ai_service.retrieve_relevant_chunks(question.text, doc_list)
        result = ai_service.generate_answer(question.text, chunks)
        Answer.objects.create(
            question=question,
            text=result['answer'],
            citations='|||'.join(result['citations']),
            confidence=result['confidence']
        )
    
    messages.success(request, 'Answers generated!')
    return redirect('review_answers', pk=questionnaire.pk)

@login_required
def review_answers_view(request, pk):
    questionnaire = get_object_or_404(Questionnaire, pk=pk, user=request.user)
    questions = questionnaire.questions.all()
    total = questions.count()
    answered = sum(1 for q in questions if hasattr(q, 'answer') and q.answer.text != 'Not found in references.')
    not_found = sum(1 for q in questions if hasattr(q, 'answer') and q.answer.text == 'Not found in references.')
    coverage = {'total': total, 'answered': answered, 'not_found': not_found}
    return render(request, 'review_answers.html', {'questionnaire': questionnaire, 'questions': questions, 'coverage': coverage})

@login_required
def update_answer_view(request, pk):
    if request.method == 'POST':
        answer = get_object_or_404(Answer, pk=pk, question__questionnaire__user=request.user)
        answer.text = request.POST.get('text')
        answer.edited = True
        answer.save()
        return JsonResponse({'success': True})
    return JsonResponse({'success': False})

@login_required
def regenerate_answer_view(request, pk):
    question = get_object_or_404(Question, pk=pk, questionnaire__user=request.user)
    documents = ReferenceDocument.objects.filter(user=request.user)
    doc_list = [{'title': doc.title, 'content': doc.content} for doc in documents]
    ai_service = AIService()
    chunks = ai_service.retrieve_relevant_chunks(question.text, doc_list)
    result = ai_service.generate_answer(question.text, chunks)
    
    if hasattr(question, 'answer'):
        answer = question.answer
        answer.text = result['answer']
        answer.citations = '|||'.join(result['citations'])
        answer.confidence = result['confidence']
        answer.save()
    else:
        Answer.objects.create(question=question, text=result['answer'], citations='|||'.join(result['citations']), confidence=result['confidence'])
    
    return JsonResponse({'success': True})

@login_required
def export_questionnaire_view(request, pk):
    questionnaire = get_object_or_404(Questionnaire, pk=pk, user=request.user)
    doc = Document()
    doc.add_heading(questionnaire.title, 0)
    
    for question in questionnaire.questions.all():
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{question.order}. {question.text}")
        q_run.bold = True
        q_run.font.size = Pt(12)
        
        if hasattr(question, 'answer'):
            answer = question.answer
            a_para = doc.add_paragraph()
            a_run = a_para.add_run(f"Answer: {answer.text}")
            a_run.font.size = Pt(11)
            
            if answer.citations:
                c_para = doc.add_paragraph()
                c_run = c_para.add_run("Citations:")
                c_run.font.size = Pt(10)
                c_run.font.color.rgb = RGBColor(100, 100, 100)
                for citation in answer.get_citations_list():
                    cite_para = doc.add_paragraph(citation, style='List Bullet')
                    cite_para.runs[0].font.size = Pt(9)
                    cite_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            
            conf_para = doc.add_paragraph()
            conf_run = conf_para.add_run(f"Confidence: {int(answer.confidence * 100)}%")
            conf_run.font.size = Pt(9)
            conf_run.font.color.rgb = RGBColor(150, 150, 150)
        else:
            a_para = doc.add_paragraph()
            a_run = a_para.add_run("Answer: Not generated yet.")
            a_run.font.size = Pt(11)
        
        doc.add_paragraph()
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{questionnaire.title}_answered.docx"'
    doc.save(response)
    return response

@login_required
def delete_questionnaire_view(request, pk):
    questionnaire = get_object_or_404(Questionnaire, pk=pk, user=request.user)
    questionnaire.delete()
    messages.success(request, 'Questionnaire deleted!')
    return redirect('dashboard')

@login_required
def delete_document_view(request, pk):
    document = get_object_or_404(ReferenceDocument, pk=pk, user=request.user)
    document.delete()
    messages.success(request, 'Document deleted!')
    return redirect('dashboard')
